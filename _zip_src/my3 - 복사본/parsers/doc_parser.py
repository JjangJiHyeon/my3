"""
Parser for .doc (binary Word 97-2003) and .docx (Open XML) files.

Strategy order for .doc:
  1. Apache Tika  (most robust, needs Java)
  2. pywin32 COM  (Windows + Word installed)
  3. OLE binary piece-table  (pure-Python fallback)
"""

from __future__ import annotations

import logging
import os
import re
import struct
from typing import Any

logger = logging.getLogger(__name__)


# ── text normalisation (shared) ──────────────────────────────────────


def _estimate_text_quality(pages: list[dict[str, Any]]) -> float:
    if not pages:
        return 0.0
    empty_pages = sum(1 for page in pages if not str(page.get("text", "") or "").strip())
    avg_text_chars = sum(len(str(page.get("text", "") or "")) for page in pages) / len(pages)
    quality = 100.0
    quality -= (empty_pages / max(1, len(pages))) * 45.0
    if avg_text_chars < 300:
        quality -= 10.0
    elif avg_text_chars > 1200:
        quality += 3.0
    return round(max(0.0, min(100.0, quality)), 2)


def _align_doc_metadata(
    metadata: dict[str, Any],
    pages: list[dict[str, Any]],
    routing_result: dict[str, Any],
    parser_subtype: str,
) -> None:
    doc_type = routing_result.get("document_type", "text_report")
    page_count = len(pages)
    empty_pages = sum(1 for page in pages if not page.get("blocks") or not str(page.get("text", "") or "").strip())
    page_type_distribution = routing_result.get("page_type_distribution") or {}
    routing_reasons = routing_result.get("routing_reasons") or [routing_result.get("routing_reason", "")]
    title_block_count = sum(
        1 for page in pages for block in page.get("blocks", []) if block.get("type") == "title"
    )
    text_block_count = sum(
        1 for page in pages for block in page.get("blocks", []) if block.get("type") == "text"
    )
    table_count = sum(len(page.get("tables", []) or []) for page in pages)

    metadata.update({
        "document_type": doc_type,
        "refined_document_type": doc_type,
        "refined_routing_type": doc_type,
        "doc_type": routing_result.get("doc_type", doc_type),
        "routing_confidence": routing_result.get("confidence", 0.0),
        "page_type_distribution": page_type_distribution,
        "routing_reasons": routing_reasons,
        "routing_signals": routing_result.get("routing_signals"),
        "routing_reason": routing_result.get("routing_reason"),
        "routing_mismatch_flag": False,
        "text_quality": _estimate_text_quality(pages),
        "empty_pages": empty_pages,
        "pipeline_used": f"doc_parser::{parser_subtype}",
        "parser_subtype": parser_subtype,
        "content_structure": {
            "page_count": page_count,
            "title_block_count": title_block_count,
            "text_block_count": text_block_count,
            "table_count": table_count,
        },
    })

def _analyze_narrative_characteristics(text: str) -> dict[str, Any]:
    import re
    clean = text.replace(" ", "").replace("\n", "")
    num_ratio = len(re.findall(r'\d', clean)) / max(1, len(clean))
    sentences = len(re.findall(r'[가-힣][다음함]\.', text))
    
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    long_paras = sum(1 for p in paragraphs if len(p) > 60 and ('다' in p or '음' in p or '함' in p))
    label_patterns = sum(1 for p in paragraphs if len(p) < 30 and re.search(r'[:\-\u2022\[\]]', p))
    
    return {
        "num_ratio": round(num_ratio, 3),
        "sentence_endings_found": sentences,
        "long_narrative_paragraphs": long_paras,
        "short_label_patterns": label_patterns,
        "is_false_positive_warning_candidate": bool(num_ratio > 0.25 and sentences <= 1 and long_paras > 0),
        "is_truly_dashboard": bool(num_ratio > 0.25 and sentences <= 1 and long_paras == 0)
    }

def _normalise_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


_WORD_FIELD_NOISE_TOKENS = {
    "MERGEFORMAT",
    "SHAPE",
    "PAGEREF",
    "HYPERLINK",
    "FORMTEXT",
    "EMBED",
    "INCLUDEPICTURE",
    "SEQ",
    "DOCVARIABLE",
    "STYLEREF",
    "TOC",
}


def _clean_ole_extracted_text(text: str) -> tuple[str, dict[str, int]]:
    cleaned_lines: list[str] = []
    removed_lines = 0
    stripped_tokens = 0

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            cleaned_lines.append("")
            continue

        upper = line.upper()
        token_hits = [token for token in _WORD_FIELD_NOISE_TOKENS if token in upper]
        if token_hits:
            line = re.sub(
                r"\b(?:MERGEFORMAT|SHAPE|PAGEREF|HYPERLINK|FORMTEXT|EMBED|INCLUDEPICTURE|SEQ|DOCVARIABLE|STYLEREF|TOC)\b",
                " ",
                line,
                flags=re.I,
            )
            line = re.sub(r"[_\\]+", " ", line)
            line = re.sub(r"\s{2,}", " ", line).strip()
            stripped_tokens += len(token_hits)

        alpha_count = len(re.findall(r"[A-Za-z가-힣]", line))
        if not line or (alpha_count == 0 and len(re.findall(r"\d", line)) < 2):
            removed_lines += 1
            continue

        if token_hits and alpha_count < 4 and len(line) < 24:
            removed_lines += 1
            continue

        cleaned_lines.append(line)

    return _normalise_text("\n".join(cleaned_lines)), {
        "word_field_noise_lines_removed": removed_lines,
        "word_field_noise_tokens_stripped": stripped_tokens,
    }


# ── section header splitting ─────────────────────────────────────────

_SECTION_HEADER_RE = re.compile(
    r'^\s*'
    r'(?:'
    r'[Ⅰ-Ⅹ]'
    r'|V?I{0,3}'
    r'|[0-9]{1,2}'
    r'|[가-힣]'
    r'|[A-Z]'
    r')'
    r'[\.\)]\s*'
    r'(.+)',
    re.MULTILINE
)
_DATE_LINE_RE = re.compile(r'^\s*\d{4}[.\-/]\d{1,2}[.\-/]\d{1,2}')
_PURE_NUMBER_RE = re.compile(r'^\s*[\d,.\-+%()△▲▽※ ]+\s*$')
_VALUE_UNIT_PATTERN = r"(?:%|억원|조원|백만원|천원|만원|원|건|배|명|개)"
_DOC_META_KEYWORDS = ("보도", "배포", "담당부서", "책임자", "담당자")


def _is_doc_section_header(line: str) -> bool:
    stripped = line.strip()
    if len(stripped) > 60:
        return False
    if _DATE_LINE_RE.match(stripped):
        return False
    if _PURE_NUMBER_RE.match(stripped):
        return False
    if _SECTION_HEADER_RE.match(stripped):
        return True
    if re.match(r'^\s*(?:[0-9]{1,2}|[Ⅰ-Ⅹ])\s+.+$', stripped):
        return True
    return False


def _is_data_row(line: str) -> bool:
    """탭 구분 숫자/금액 행인지 판단 (표 데이터 행)"""
    stripped = line.strip()
    if not stripped:
        return False
    # 탭이 여러 개 있고 숫자 비율이 높은 행
    if '\t' in stripped:
        parts = [p.strip() for p in stripped.split('\t') if p.strip()]
        if len(parts) >= 2:
            num_parts = sum(1 for p in parts if _PURE_NUMBER_RE.match(p))
            if num_parts >= len(parts) * 0.5:
                return True
    # 순수 숫자/기호 행
    if _PURE_NUMBER_RE.match(stripped):
        return True
    return False


def _split_tabular_cells(line: str) -> list[str]:
    stripped = line.strip()
    if not stripped:
        return []
    if "|" in stripped:
        cells = [cell.strip() for cell in stripped.split("|") if cell.strip()]
        if len(cells) >= 2:
            return cells
    cells = [cell.strip() for cell in re.split(r"\t+|\s{2,}", stripped) if cell.strip()]
    if len(cells) >= 2:
        return cells
    tokens = stripped.split()
    if len(tokens) >= 4:
        paired_cells: list[str] = []
        label_tokens: list[str] = []
        value_pairs = 0
        for token in tokens:
            is_value = bool(re.fullmatch(r"[\d,]+(?:\.\d+)?(?:%|억원|조원|백만원|천원|만원|원|건|배)?", token))
            if is_value and label_tokens:
                paired_cells.append(" ".join(label_tokens).strip())
                paired_cells.append(token)
                label_tokens.clear()
                value_pairs += 1
            else:
                label_tokens.append(token)
        if value_pairs >= 2 and len(paired_cells) >= 4:
            return paired_cells
    summary_pair = _split_summary_label_value(stripped)
    if summary_pair:
        return [summary_pair[0], summary_pair[1]]
    return []


def _looks_value_like(text: str) -> bool:
    compact = " ".join(text.split())
    if len(compact) > 42 or not re.search(r"\d", compact):
        return False
    if re.search(r"(감소|증가|발행|조달|대비|으로|이며|합니다|하였다|했다)", compact):
        return False
    unit_hits = len(re.findall(_VALUE_UNIT_PATTERN, compact))
    alpha_groups = len(re.findall(r"[A-Za-z가-힣]+", compact))
    return bool(unit_hits or len(re.findall(r"\d", compact)) >= 3) and alpha_groups <= 2


def _split_summary_label_value(text: str) -> tuple[str, str] | None:
    stripped = text.strip()
    if not stripped:
        return None
    if stripped.startswith("("):
        return None
    if re.match(r'^\s*([□◦▪•\-\u25e6\u2022\u25a1\u25cb\u2023\u25a0])', stripped):
        return None
    tokens = stripped.split()
    if len(tokens) < 2:
        return None
    for idx in range(1, len(tokens)):
        label = " ".join(tokens[:idx]).strip()
        value = " ".join(tokens[idx:]).strip()
        if len(tokens[idx:]) > 3:
            continue
        if not label or len(label) > 42 or _is_list_marker(label):
            continue
        if re.search(r"\d{2,}", label):
            continue
        if _looks_value_like(value):
            return label, value
    return None


def _split_doc_line_fragments(line: str) -> list[str]:
    working = (
        line.replace("\u2005", " ")
        .replace("\u2002", " ")
        .replace("\u2003", " ")
        .replace("\u2009", " ")
        .replace("\xa0", " ")
    )
    working = re.sub(r"\s{3,}(?=(?:전월 대비|['’]?\d{2,4}[.\-/]|[0-9]{1,2}\s+[가-힣]|보도|배포|담당부서|책임자|담당자))", "\n", working)
    working = re.sub(r"(?<!^)(보도|배포|담당부서|책임자|담당자)(?=[가-힣A-Za-z0-9(])", r"\n\1", working)
    working = re.sub(r"(보도|배포|담당부서|책임자|담당자)\s*(?=[가-힣0-9'’(])", r"\1\n", working)
    working = re.sub(r"(\(\d{2,4}-\d{3,4}-\d{4}\))(?=[가-힣A-Za-z'’])", r"\1\n", working)
    working = re.sub(r"(?<=\))\s+(?=(?:[0-9]{1,2}|[Ⅰ-Ⅹ]|[가-힣])(?:\s+|[.)]))", "\n", working)
    fragments: list[str] = []
    for piece in working.splitlines():
        stripped = " ".join(piece.split())
        if not stripped:
            fragments.append("")
            continue
        note_match = re.match(r"^(\([^)]{4,120}\))\s*(.*)$", stripped)
        if note_match:
            note = note_match.group(1).strip()
            rest = note_match.group(2).strip()
            fragments.append(note)
            if rest:
                summary_pair = _split_summary_label_value(rest)
                if summary_pair:
                    fragments.append(f"{summary_pair[0]}\t{summary_pair[1]}")
                else:
                    fragments.append(rest)
            continue

        trailing_note = re.match(r"^(.+?)\s+(\([^)]{4,120}\))\s*$", stripped)
        if trailing_note:
            main_part = trailing_note.group(1).strip()
            note_part = trailing_note.group(2).strip()
            summary_pair = _split_summary_label_value(main_part)
            if summary_pair:
                fragments.append(f"{summary_pair[0]}\t{summary_pair[1]}")
            else:
                fragments.append(main_part)
            fragments.append(note_part)
            continue

        summary_pair = _split_summary_label_value(stripped)
        if summary_pair and len(stripped) <= 48 and stripped.count(" ") <= 4:
            fragments.append(f"{summary_pair[0]}\t{summary_pair[1]}")
            continue
        fragments.append(stripped)
    return fragments


def _looks_doc_meta_line(line: str) -> bool:
    stripped = line.strip().lstrip("*").strip()
    if not stripped:
        return False
    if any(stripped.startswith(keyword) for keyword in _DOC_META_KEYWORDS):
        return True
    if re.search(r"\(\d{2,4}-\d{3,4}-\d{4}\)", stripped):
        return True
    if _DATE_LINE_RE.match(stripped):
        return True
    if re.match(r"^['’]?\d{2,4}\.[0-9]{1,2}월", stripped):
        return True
    return False


def _looks_doc_display_title(line: str) -> bool:
    stripped = line.strip().lstrip("*").strip()
    if not stripped or _looks_doc_meta_line(stripped):
        return False
    if re.search(r"\(\d{2,4}-\d{3,4}-\d{4}\)", stripped):
        return False
    if re.fullmatch(r"(?:제\s*목|보\s*도\s*자\s*료)", stripped):
        return True
    if re.match(r"^['’]?\d{2,4}\.[0-9]{1,2}월", stripped) and len(stripped) <= 40:
        return True
    if re.fullmatch(r"[가-힣A-Za-z0-9](?:\s+[가-힣A-Za-z0-9]){2,15}", stripped) and len(stripped.replace(" ", "")) <= 16:
        return True
    return False


def _is_list_marker(cell: str) -> bool:
    return bool(re.fullmatch(r"[□◦▪•\-\u25e6\u2022\u25a1\u25cb\u2023\u25a0]|[\d]{1,2}[\.\)]|[가-힣][\.\)]", cell.strip()))


def _row_noise_token_ratio(row: list[str]) -> float:
    if not row:
        return 0.0
    noisy = 0
    for cell in row:
        upper = cell.upper()
        if any(token in upper for token in _WORD_FIELD_NOISE_TOKENS):
            noisy += 1
    return noisy / len(row)


def _is_compact_label_value_row(row: list[str]) -> bool:
    if len(row) < 4 or len(row) % 2 != 0:
        return False
    pair_count = 0
    for i in range(0, len(row), 2):
        label = row[i].strip()
        value = row[i + 1].strip()
        if not label or not value:
            return False
        if len(label) > 30:
            return False
        if not re.search(r"\d", value):
            return False
        pair_count += 1
    return pair_count >= 2


def _is_summary_pair_row(row: list[str]) -> bool:
    if len(row) != 2:
        return False
    label = row[0].strip()
    value = row[1].strip()
    if not label or not value or len(label) > 42 or _is_list_marker(label):
        return False
    if re.search(r"\d{2,}", label):
        return False
    return _looks_value_like(value)


def _should_promote_table_rows(table_rows: list[list[str]]) -> tuple[bool, str]:
    if not table_rows:
        return False, "empty_rows"

    row_lengths = [len(row) for row in table_rows if row]
    if not row_lengths:
        return False, "empty_rows"

    shape_cols = max(row_lengths)
    stable_rows = sum(1 for length in row_lengths if abs(length - shape_cols) <= 1)
    stability = stable_rows / len(row_lengths)
    numeric_cells = sum(1 for row in table_rows for cell in row if re.search(r"\d", cell))
    total_cells = sum(len(row) for row in table_rows)
    numeric_ratio = numeric_cells / max(1, total_cells)
    noise_ratio = sum(_row_noise_token_ratio(row) for row in table_rows) / len(table_rows)

    if len(table_rows) == 1:
        if _is_compact_label_value_row(table_rows[0]):
            return True, "compact_label_value_row"
        if _is_summary_pair_row(table_rows[0]):
            return True, "single_summary_pair_row"
        return False, "single_row_not_compact"

    if shape_cols < 2 or stability < 0.7:
        return False, "unstable_columns"
    if noise_ratio > 0.25:
        return False, "word_field_noise"

    first_col = [row[0].strip() for row in table_rows if row]
    first_col_list_ratio = sum(1 for cell in first_col if _is_list_marker(cell)) / max(1, len(first_col))
    second_col_avg_len = sum(len(row[1].strip()) for row in table_rows if len(row) > 1) / max(1, sum(1 for row in table_rows if len(row) > 1))
    if shape_cols == 2 and first_col_list_ratio > 0.6 and second_col_avg_len > 35:
        return False, "bullet_narrative_list"

    long_cell_ratio = sum(1 for row in table_rows for cell in row if len(cell) > 70) / max(1, total_cells)
    if long_cell_ratio > 0.35 and numeric_ratio < 0.2:
        return False, "paragraph_like_cells"

    first_row = table_rows[0]
    header_like = sum(1 for cell in first_row if not re.search(r"\d", cell) and len(cell) <= 30) / max(1, len(first_row))
    repeated_label_value = sum(1 for row in table_rows if _is_compact_label_value_row(row)) >= max(1, len(table_rows) // 2)
    unit_ratio = sum(
        1
        for row in table_rows
        for cell in row
        if re.search(r"(억원|조원|백만원|천원|%|건|배|원)$", cell.strip())
    ) / max(1, total_cells)
    summary_pair_rows = 0
    for row in table_rows:
        if _is_summary_pair_row(row):
            summary_pair_rows += 1
    if len(table_rows) >= 2 and summary_pair_rows >= max(2, len(table_rows) - 1):
        return True, "summary_pair_rows"
    if numeric_ratio >= 0.18 or header_like >= 0.5 or repeated_label_value:
        return True, "stable_tabular_pattern"
    return False, "weak_tabular_signal"


def _looks_table_like_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if len(stripped) > 120 and "|" not in stripped and "\t" not in stripped:
        return False
    cells = _split_tabular_cells(line)
    if len(cells) < 2:
        return False
    numeric_like = sum(1 for cell in cells if re.search(r"\d", cell))
    compact_like = sum(1 for cell in cells if len(cell) <= 30)
    if _row_noise_token_ratio(cells) > 0.3:
        return False
    if len(cells) == 2 and _is_list_marker(cells[0]) and len(cells[1]) > 35:
        return False
    return (
        numeric_like >= 1
        or compact_like >= max(2, len(cells) - 1)
        or _is_compact_label_value_row(cells)
        or _is_summary_pair_row(cells)
    )


def _table_text_from_rows(rows: list[list[str]]) -> str:
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows if any(cell.strip() for cell in row))


def _flush_table_block(
    blocks: list[dict[str, Any]],
    table_rows: list[list[str]],
    page_num: int,
    block_idx: int,
    source: str,
) -> int:
    should_promote, structure_reason = _should_promote_table_rows(table_rows)
    if not should_promote:
        return block_idx
    shape_cols = max((len(row) for row in table_rows), default=0)
    normalized_rows = [row + [""] * max(0, shape_cols - len(row)) for row in table_rows]
    blocks.append({
        "id": f"p{page_num}_b{block_idx}",
        "type": "table",
        "bbox": [0, 0, 0, 0],
        "text": _table_text_from_rows(normalized_rows),
        "page_num": page_num,
        "source": source,
        "score": 1.0,
        "meta": {
            "rows": normalized_rows,
            "block_subtype": "table_like_rows",
            "table_candidate_score": 3.0,
            "table_shape": {"rows": len(normalized_rows), "cols": shape_cols},
            "structure_reason": structure_reason,
        }
    })
    return block_idx + 1


def _structure_text_blocks(
    text: str,
    page_num: int,
    source: str,
) -> tuple[list[dict[str, Any]], list[list[list[str]]], dict[str, int]]:
    if not text or not text.strip():
        return [], [], {"title_count": 0, "text_count": 0, "table_count": 0, "table_like_line_count": 0, "meta_info_count": 0}
    lines: list[str] = []
    for raw_line in text.split('\n'):
        lines.extend(_split_doc_line_fragments(raw_line))
    blocks: list[dict[str, Any]] = []
    tables: list[list[list[str]]] = []
    block_idx = 0
    current_body_lines: list[str] = []
    current_data_lines: list[str] = []
    current_table_rows: list[list[str]] = []
    pending_meta_value_lines = 0
    meta_accumulator: list[str] = []
    stats = {"title_count": 0, "text_count": 0, "table_count": 0, "table_like_line_count": 0, "meta_info_count": 0}

    def _flush_body():
        nonlocal block_idx
        body_text = '\n'.join(current_body_lines).strip()
        if body_text:
            blocks.append({
                "id": f"p{page_num}_b{block_idx}",
                "type": "text",
                "bbox": [0, 0, 0, 0],
                "text": body_text,
                "page_num": page_num,
                "source": source,
                "score": 1.0,
                "meta": {}
            })
            block_idx += 1
            stats["text_count"] += 1
        current_body_lines.clear()

    def _flush_data():
        nonlocal block_idx
        data_text = '\n'.join(current_data_lines).strip()
        if data_text:
            blocks.append({
                "id": f"p{page_num}_b{block_idx}",
                "type": "text",
                "bbox": [0, 0, 0, 0],
                "text": data_text,
                "page_num": page_num,
                "source": source,
                "score": 1.0,
                "meta": {"block_subtype": "data_rows"}
            })
            block_idx += 1
            stats["text_count"] += 1
        current_data_lines.clear()

    def _flush_table():
        nonlocal block_idx
        promoted_before = block_idx
        if current_table_rows:
            block_idx = _flush_table_block(blocks, current_table_rows, page_num, block_idx, source)
        if block_idx > promoted_before:
            tables.append([row[:] for row in current_table_rows])
            stats["table_count"] += 1
        elif current_table_rows:
            current_body_lines.extend(" ".join(row) for row in current_table_rows)
        current_table_rows.clear()

    def _flush_meta():
        nonlocal block_idx
        if not meta_accumulator:
            return
        meta_text = "\n".join(meta_accumulator).strip()
        if meta_text:
            blocks.append({
                "id": f"p{page_num}_b{block_idx}",
                "type": "text",
                "bbox": [0, 0, 0, 0],
                "text": meta_text,
                "page_num": page_num,
                "source": source,
                "score": 1.0,
                "meta": {"block_subtype": "meta_info_group"}
            })
            block_idx += 1
            stats["text_count"] += 1
            stats["meta_info_count"] += 1
        meta_accumulator.clear()

    for line in lines:
        stripped = line.strip()
        is_header = _is_doc_section_header(line) or _looks_doc_display_title(line)
        is_data = _is_data_row(line)
        is_table_like = _looks_table_like_line(line)
        # 목록 기호 시작 판정 (□, ◦, -, 1),   ◦ 등)
        is_list_item = len(stripped) < 200 and re.match(r'^\s*([□◦\-\u25e6\u2022\u25a1\u25cb\u2023\u25a0]|[\d]{1,2}[\.\)])', line)
        meta_label_line = any(stripped.startswith(keyword) for keyword in _DOC_META_KEYWORDS)
        inherited_meta = (
            pending_meta_value_lines > 0
            and stripped
            and len(stripped) <= 80
            and not is_header
            and not is_table_like
            and not is_data
            and not is_list_item
        )
        is_meta_info = (not is_header) and (_looks_doc_meta_line(line) or inherited_meta)

        if stripped and not is_meta_info and not meta_label_line:
            pending_meta_value_lines = 0

        if is_header:
            _flush_table()
            _flush_data()
            _flush_body()
            _flush_meta()
            blocks.append({
                "id": f"p{page_num}_b{block_idx}",
                "type": "title",
                "bbox": [0, 0, 0, 0],
                "text": line.strip().lstrip("*").strip(),
                "page_num": page_num,
                "source": source,
                "score": 1.0,
                "meta": {"block_subtype": "document_title" if _looks_doc_display_title(line) else "section_title"}
            })
            block_idx += 1
            stats["title_count"] += 1
        elif is_meta_info:
            _flush_table()
            _flush_data()
            _flush_body()
            meta_accumulator.append(line.strip().lstrip("*").strip())
            if meta_label_line:
                pending_meta_value_lines = 2
            elif inherited_meta:
                pending_meta_value_lines = max(0, pending_meta_value_lines - 1)
        elif not stripped:
            _flush_table()
            _flush_data()
            _flush_body()
            _flush_meta()
            pending_meta_value_lines = 0
        elif is_table_like:
            _flush_body()
            _flush_data()
            _flush_meta()
            current_table_rows.append(_split_tabular_cells(line))
            stats["table_like_line_count"] += 1
        elif is_list_item:
            _flush_table()
            _flush_data()
            _flush_body()
            _flush_meta()
            blocks.append({
                "id": f"p{page_num}_b{block_idx}",
                "type": "text",
                "bbox": [0, 0, 0, 0],
                "text": line.strip(),
                "page_num": page_num,
                "source": source,
                "score": 1.0,
                "meta": {"block_subtype": "list_item"}
            })
            block_idx += 1
            stats["text_count"] += 1
        elif is_data:
            # 데이터 행: 바디 버퍼가 있으면 비우고 데이터 버퍼로
            _flush_table()
            if current_body_lines:
                _flush_body()
            current_data_lines.append(line)
        else:
            # 일반 텍스트: 데이터 버퍼가 있으면 비우고 바디 버퍼로
            _flush_table()
            if current_data_lines:
                _flush_data()
            current_body_lines.append(line)

    _flush_table()
    _flush_data()
    _flush_body()
    _flush_meta()
    
    # 흡수 최적화 (파편 문장 합치기)
    blocks = _absorb_tiny_fragments_sequential(blocks)
    
    stats["title_count"] = sum(1 for block in blocks if block["type"] == "title")
    stats["text_count"] = sum(1 for block in blocks if block["type"] == "text")
    return blocks, tables, stats


def _split_into_section_blocks(text: str, page_num: int = 1) -> list[dict[str, Any]]:
    blocks, _, _ = _structure_text_blocks(text, page_num, "doc_parser")
    return blocks

def _absorb_tiny_fragments_sequential(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """DOC/HWP 등 bbox가 없는 순차적 블록의 조각을 인접 블록으로 흡수"""
    if not blocks: return []
    out = []
    
    for i, b in enumerate(blocks):
        text = b.get("text", "")
        # Tiny chunk (under 30 chars, single line, no structural header numbers)
        if b["type"] == "text" and len(text) < 30 and "\n" not in text and not (len(text) < 5 and text.strip().isdigit()):
            # Absorb to the previous block if text type, otherwise next block
            if out and out[-1]["type"] == "text":
                out[-1]["text"] += "\n" + text.strip()
            elif i + 1 < len(blocks) and blocks[i+1]["type"] == "text":
                blocks[i+1]["text"] = text.strip() + "\n" + blocks[i+1].get("text", "")
            else:
                out.append(b)
        else:
            out.append(b)
            
    return out


def _generate_doc_rag_text(blocks: list[dict[str, Any]]) -> str:
    """DOC 전용 RAG 텍스트 생성 (PDF와 유사한 스키마 사용)"""
    parts = []
    for b in blocks:
        btype = b["type"]
        text = b.get("text", "").strip()
        if not text:
            continue
        meta = b.get("meta", {})
        subtype = meta.get("block_subtype")

        if btype == "title":
            parts.append(f"\n[SECTION: {text}]\n")
        elif btype == "text":
            if subtype in ("meta_info", "meta_info_group"):
                parts.append(f"[META] {text}")
            else:
                parts.append(text)
        elif btype == "table":
            rows = meta.get("rows", [])
            row_count = len(rows)
            col_count = max((len(row) for row in rows), default=0)
            if row_count <= 20:
                parts.append(f"\n[TABLE: {row_count} rows, {col_count} cols]\n{text}\n")
            else:
                parts.append(f"\n[TABLE: {row_count} rows, {col_count} cols]\n{text[:400]}\n")
    return "\n\n".join(parts).strip()


# ── public entry point ───────────────────────────────────────────────

def parse_doc(filepath: str) -> dict[str, Any]:
    ext = os.path.splitext(filepath)[1].lower()
    errors: list[str] = []

    strategies = []
    if ext == ".docx":
        strategies.append(("python-docx", _parse_docx))
    else:
        strategies.extend([
            ("tika", _parse_doc_tika),
            ("win32com", _parse_doc_win32),
            ("OLE binary", _parse_doc_binary),
        ])

    for name, fn in strategies:
        try:
            result = fn(filepath)
            result.setdefault("metadata", {})["parse_strategies_tried"] = (
                errors if errors else [f"{name} (first attempt)"]
            )
            
            # ── Document Routing ──────────────────────────────
            from .document_router import route_document
            pages = result.get("pages", [])
            meta = result.get("metadata", {})
            routing_result = route_document(None, pages, metadata=meta)
            
            doc_type = routing_result.get("document_type", "text_report")
            _align_doc_metadata(meta, pages, routing_result, parser_subtype=str(meta.get("parser_used", name)))
            
            for p in pages:
                dbg = p.setdefault("parser_debug", {})
                dbg["document_type"] = doc_type
                dbg["refined_document_type"] = doc_type
                dbg["doc_type"] = doc_type
                dbg["pipeline_used"] = meta.get("pipeline_used")
                # RAG Text Generation
                p["rag_text"] = _generate_doc_rag_text(p.get("blocks", []))
                
            return result
        except Exception as exc:
            msg = f"{name}: {exc}"
            errors.append(msg)
            logger.info("doc strategy '%s' failed: %s", name, exc)

    return {
        "pages": [],
        "metadata": {"parse_strategies_tried": errors},
        "status": "error",
        "error": "All parsing strategies failed: " + " | ".join(errors),
    }


# ── Strategy: python-docx (.docx) ───────────────────────────────────

def _parse_docx(filepath: str) -> dict[str, Any]:
    from docx import Document

    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = _normalise_text("\n".join(paragraphs))

    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    metadata: dict[str, Any] = {
        "page_count": 1,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "char_count": len(full_text),
        "parser_used": "python-docx",
    }

    sec_blocks, inferred_tables, structure_stats = _structure_text_blocks(full_text, 1, "doc_parser")
    all_tables = tables + [table for table in inferred_tables if table]
    pages = [{
        "page_num": 1,
        "page_width": 0,
        "page_height": 0,
        "preview_width": 0,
        "preview_height": 0,
        "preview_scale_x": 1.0,
        "preview_scale_y": 1.0,
        "coord_space": "page_points",
        "preview_image": None,
        "text": full_text,
        "tables": all_tables,
        "blocks": sec_blocks,
        "image_count": 0,
        "text_source": "native",
        "ocr_applied": False,
        "ocr_confidence": 0.0,
        "parser_debug": {
            "preview_generated": False,
            "preview_error": None,
            "native_text_chars": len(full_text),
            "ocr_used": False,
            "ocr_trigger_reason": "ocr_not_needed",
            "candidate_counts": {
                "raw_text_blocks": 1 if full_text else 0,
                "final_blocks": len(sec_blocks)
            },
            "block_type_counts": {
                "title": sum(1 for b in sec_blocks if b["type"] == "title"),
                "text": sum(1 for b in sec_blocks if b["type"] == "text"),
                "table": sum(1 for b in sec_blocks if b["type"] == "table"),
            },
            "dropped_blocks": [],
            "bbox_warnings": [],
            "narrative_analysis": _analyze_narrative_characteristics(full_text),
            "structure_extraction": structure_stats,
        }
    }]
    
    metadata["narrative_analysis"] = pages[0]["parser_debug"]["narrative_analysis"]
    return {"pages": pages, "metadata": metadata, "status": "success"}


# ── Strategy 1: Tika ────────────────────────────────────────────────

def _parse_doc_tika(filepath: str) -> dict[str, Any]:
    try:
        from tika import parser as tika_parser
    except ImportError:
        raise RuntimeError("tika package not installed")

    parsed = tika_parser.from_file(filepath)
    content = parsed.get("content")
    if not content or not content.strip():
        raise ValueError("Tika returned empty content")

    full_text = _normalise_text(content)
    tika_meta = parsed.get("metadata") or {}

    metadata: dict[str, Any] = {
        "page_count": int(tika_meta.get("xmpTPg:NPages", 1) or 1),
        "char_count": len(full_text),
        "parser_used": "Apache Tika",
        "tika_content_type": tika_meta.get("Content-Type", ""),
    }

    sec_blocks, inferred_tables, structure_stats = _structure_text_blocks(full_text, 1, "doc_parser")
    pages = [{
        "page_num": 1,
        "page_width": 0,
        "page_height": 0,
        "preview_width": 0,
        "preview_height": 0,
        "preview_scale_x": 1.0,
        "preview_scale_y": 1.0,
        "coord_space": "page_points",
        "preview_image": None,
        "text": full_text,
        "tables": inferred_tables,
        "blocks": sec_blocks,
        "image_count": 0,
        "text_source": "native",
        "ocr_applied": False,
        "ocr_confidence": 0.0,
        "parser_debug": {
            "preview_generated": False,
            "preview_error": None,
            "native_text_chars": len(full_text),
            "ocr_used": False,
            "ocr_trigger_reason": "ocr_not_needed",
            "candidate_counts": {
                "raw_text_blocks": 1 if full_text else 0,
                "final_blocks": len(sec_blocks)
            },
            "block_type_counts": {
                "title": sum(1 for b in sec_blocks if b["type"] == "title"),
                "text": sum(1 for b in sec_blocks if b["type"] == "text"),
                "table": sum(1 for b in sec_blocks if b["type"] == "table"),
            },
            "dropped_blocks": [],
            "bbox_warnings": [],
            "narrative_analysis": _analyze_narrative_characteristics(full_text),
            "structure_extraction": structure_stats,
        }
    }]
    
    metadata["narrative_analysis"] = pages[0]["parser_debug"]["narrative_analysis"]
    return {"pages": pages, "metadata": metadata, "status": "success"}


# ── Strategy 2: pywin32 COM ─────────────────────────────────────────

def _parse_doc_win32(filepath: str) -> dict[str, Any]:
    try:
        import win32com.client  # type: ignore
    except ImportError:
        raise RuntimeError(
            "pywin32 is not installed. Install with: pip install pywin32  "
            "(requires Microsoft Word on this machine)"
        )

    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(filepath), ReadOnly=True)
        full_text = _normalise_text(doc.Content.Text)
        doc.Close(False)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass

    metadata: dict[str, Any] = {
        "page_count": 1,
        "char_count": len(full_text),
        "parser_used": "win32com (Microsoft Word)",
    }
    sec_blocks, inferred_tables, structure_stats = _structure_text_blocks(full_text, 1, "doc_parser")
    pages = [{
        "page_num": 1,
        "page_width": 0,
        "page_height": 0,
        "preview_width": 0,
        "preview_height": 0,
        "preview_scale_x": 1.0,
        "preview_scale_y": 1.0,
        "coord_space": "page_points",
        "preview_image": None,
        "text": full_text,
        "tables": inferred_tables,
        "blocks": sec_blocks,
        "image_count": 0,
        "text_source": "native",
        "ocr_applied": False,
        "ocr_confidence": 0.0,
        "parser_debug": {
            "preview_generated": False,
            "preview_error": None,
            "native_text_chars": len(full_text),
            "ocr_used": False,
            "ocr_trigger_reason": "ocr_not_needed",
            "candidate_counts": {
                "raw_text_blocks": 1 if full_text else 0,
                "final_blocks": len(sec_blocks)
            },
            "block_type_counts": {
                "title": sum(1 for b in sec_blocks if b["type"] == "title"),
                "text": sum(1 for b in sec_blocks if b["type"] == "text"),
                "table": sum(1 for b in sec_blocks if b["type"] == "table"),
            },
            "dropped_blocks": [],
            "bbox_warnings": [],
            "narrative_analysis": _analyze_narrative_characteristics(full_text),
            "structure_extraction": structure_stats,
        }
    }]
    
    metadata["narrative_analysis"] = pages[0]["parser_debug"]["narrative_analysis"]
    return {"pages": pages, "metadata": metadata, "status": "success"}


# ── Strategy 3: OLE binary piece-table ──────────────────────────────

def _parse_doc_binary(filepath: str) -> dict[str, Any]:
    import olefile

    if not olefile.isOleFile(filepath):
        raise ValueError("Not an OLE2 compound file")

    ole = olefile.OleFileIO(filepath)
    try:
        full_text = _extract_ole_text(ole)
    finally:
        ole.close()

    full_text, ole_cleanup_stats = _clean_ole_extracted_text(full_text)

    metadata: dict[str, Any] = {
        "page_count": 1,
        "char_count": len(full_text),
        "parser_used": "OLE binary (piece table)",
    }
    sec_blocks, inferred_tables, structure_stats = _structure_text_blocks(full_text, 1, "doc_parser")
    pages = [{
        "page_num": 1,
        "page_width": 0,
        "page_height": 0,
        "preview_width": 0,
        "preview_height": 0,
        "preview_scale_x": 1.0,
        "preview_scale_y": 1.0,
        "coord_space": "page_points",
        "preview_image": None,
        "text": full_text,
        "tables": inferred_tables,
        "blocks": sec_blocks,
        "image_count": 0,
        "text_source": "native",
        "ocr_applied": False,
        "ocr_confidence": 0.0,
        "parser_debug": {
            "preview_generated": False,
            "preview_error": None,
            "native_text_chars": len(full_text),
            "ocr_used": False,
            "ocr_trigger_reason": "ocr_not_needed",
            "candidate_counts": {
                "raw_text_blocks": 1 if full_text else 0,
                "final_blocks": len(sec_blocks)
            },
            "block_type_counts": {
                "title": sum(1 for b in sec_blocks if b["type"] == "title"),
                "text": sum(1 for b in sec_blocks if b["type"] == "text"),
                "table": sum(1 for b in sec_blocks if b["type"] == "table"),
            },
            "dropped_blocks": [],
            "bbox_warnings": [],
            "narrative_analysis": _analyze_narrative_characteristics(full_text),
            "structure_extraction": {
                **structure_stats,
                **ole_cleanup_stats,
            },
        }
    }]
    
    metadata["narrative_analysis"] = pages[0]["parser_debug"]["narrative_analysis"]
    return {"pages": pages, "metadata": metadata, "status": "success"}


def _extract_ole_text(ole: Any) -> str:
    word_doc = ole.openstream("WordDocument").read()

    magic = struct.unpack_from("<H", word_doc, 0)[0]
    if magic not in (0xA5EC, 0xA5DC):
        raise ValueError(f"Invalid Word magic: {hex(magic)}")

    flags = struct.unpack_from("<H", word_doc, 0x000A)[0]
    table_name = "1Table" if (flags & 0x0200) else "0Table"

    if not ole.exists(table_name):
        alt = "0Table" if table_name == "1Table" else "1Table"
        if ole.exists(alt):
            table_name = alt
        else:
            raise ValueError("No Table stream found")

    table_stream = ole.openstream(table_name).read()

    csw = struct.unpack_from("<H", word_doc, 0x20)[0]
    fibrg_w_end = 0x22 + csw * 2

    cslw = struct.unpack_from("<H", word_doc, fibrg_w_end)[0]
    fibrg_lw_start = fibrg_w_end + 2
    fibrg_lw_end = fibrg_lw_start + cslw * 4

    ccpText = struct.unpack_from("<i", word_doc, fibrg_lw_start)[0]

    cb_rg = struct.unpack_from("<H", word_doc, fibrg_lw_end)[0]
    fclcb_start = fibrg_lw_end + 2

    if cb_rg < 34:
        raise ValueError(f"FIBRgFcLcb too small ({cb_rg} pairs)")

    fc_clx = struct.unpack_from("<I", word_doc, fclcb_start + 33 * 8)[0]
    lcb_clx = struct.unpack_from("<I", word_doc, fclcb_start + 33 * 8 + 4)[0]

    if lcb_clx == 0:
        raise ValueError("CLX has zero length")

    clx = table_stream[fc_clx: fc_clx + lcb_clx]

    off = 0
    while off < len(clx) and clx[off] == 0x01:
        cb = struct.unpack_from("<H", clx, off + 1)[0]
        off += 3 + cb

    if off >= len(clx) or clx[off] != 0x02:
        raise ValueError("Pcdt marker (0x02) not found")

    off += 1
    lcb_pcdt = struct.unpack_from("<I", clx, off)[0]
    off += 4
    plc = clx[off: off + lcb_pcdt]

    n_pieces = (lcb_pcdt - 4) // 12
    if n_pieces <= 0:
        raise ValueError("No text pieces found")

    cps = [struct.unpack_from("<I", plc, i * 4)[0] for i in range(n_pieces + 1)]
    pcd_base = (n_pieces + 1) * 4

    parts: list[str] = []
    for i in range(n_pieces):
        cp_start, cp_end = cps[i], cps[i + 1]
        if cp_start >= ccpText:
            break
        n_chars = min(cp_end, ccpText) - cp_start

        pcd = plc[pcd_base + i * 8: pcd_base + (i + 1) * 8]
        fc_raw = struct.unpack_from("<I", pcd, 2)[0]
        compressed = bool(fc_raw & (1 << 30))
        fc_val = fc_raw & 0x3FFFFFFF

        if compressed:
            boff = fc_val // 2
            parts.append(word_doc[boff: boff + n_chars].decode("cp1252", errors="replace"))
        else:
            boff = fc_val
            parts.append(word_doc[boff: boff + n_chars * 2].decode("utf-16-le", errors="replace"))

    return "".join(parts)
