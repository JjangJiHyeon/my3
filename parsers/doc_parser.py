"""
Parser for .doc (binary Word 97-2003) and .docx (Open XML) files.

Strategy order for .doc:
  1. Apache Tika  (most robust, needs Java)
  2. pywin32 COM  (Windows + Word installed)
  3. OLE binary piece-table  (pure-Python fallback)
"""

from __future__ import annotations

import logging
import os
import re
import struct
from typing import Any

logger = logging.getLogger(__name__)


# ── text normalisation (shared) ──────────────────────────────────────

def _analyze_narrative_characteristics(text: str) -> dict[str, Any]:
    import re
    clean = text.replace(" ", "").replace("\n", "")
    num_ratio = len(re.findall(r'\d', clean)) / max(1, len(clean))
    sentences = len(re.findall(r'[가-힣][다음함]\.', text))
    
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    long_paras = sum(1 for p in paragraphs if len(p) > 60 and ('다' in p or '음' in p or '함' in p))
    label_patterns = sum(1 for p in paragraphs if len(p) < 30 and re.search(r'[:\-\u2022\[\]]', p))
    
    return {
        "num_ratio": round(num_ratio, 3),
        "sentence_endings_found": sentences,
        "long_narrative_paragraphs": long_paras,
        "short_label_patterns": label_patterns,
        "is_false_positive_warning_candidate": bool(num_ratio > 0.25 and sentences <= 1 and long_paras > 0),
        "is_truly_dashboard": bool(num_ratio > 0.25 and sentences <= 1 and long_paras == 0)
    }

def _normalise_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── section header splitting ─────────────────────────────────────────

_SECTION_HEADER_RE = re.compile(
    r'^\s*'
    r'(?:'
    r'[Ⅰ-Ⅹ]'
    r'|V?I{0,3}'
    r'|[0-9]{1,2}'
    r'|[가-힣]'
    r'|[A-Z]'
    r')'
    r'[\.\)]\s*'
    r'(.+)',
    re.MULTILINE
)
_DATE_LINE_RE = re.compile(r'^\s*\d{4}[.\-/]\d{1,2}[.\-/]\d{1,2}')
_PURE_NUMBER_RE = re.compile(r'^\s*[\d,.\-+%()△▲▽※ ]+\s*$')


def _is_doc_section_header(line: str) -> bool:
    stripped = line.strip()
    if len(stripped) > 60:
        return False
    if _DATE_LINE_RE.match(stripped):
        return False
    if _PURE_NUMBER_RE.match(stripped):
        return False
    if _SECTION_HEADER_RE.match(stripped):
        return True
    return False


def _is_data_row(line: str) -> bool:
    """탭 구분 숫자/금액 행인지 판단 (표 데이터 행)"""
    stripped = line.strip()
    if not stripped:
        return False
    # 탭이 여러 개 있고 숫자 비율이 높은 행
    if '\t' in stripped:
        parts = [p.strip() for p in stripped.split('\t') if p.strip()]
        if len(parts) >= 2:
            num_parts = sum(1 for p in parts if _PURE_NUMBER_RE.match(p))
            if num_parts >= len(parts) * 0.5:
                return True
    # 순수 숫자/기호 행
    if _PURE_NUMBER_RE.match(stripped):
        return True
    return False


def _split_into_section_blocks(text: str, page_num: int = 1) -> list[dict[str, Any]]:
    if not text or not text.strip():
        return []
    lines = text.split('\n')
    blocks: list[dict[str, Any]] = []
    block_idx = 0
    current_body_lines: list[str] = []
    current_data_lines: list[str] = []

    def _flush_body():
        nonlocal block_idx
        body_text = '\n'.join(current_body_lines).strip()
        if body_text:
            blocks.append({
                "id": f"p{page_num}_b{block_idx}",
                "type": "text",
                "bbox": [0, 0, 0, 0],
                "text": body_text,
                "page_num": page_num,
                "source": "doc_parser",
                "score": 1.0,
                "meta": {}
            })
            block_idx += 1
        current_body_lines.clear()

    def _flush_data():
        nonlocal block_idx
        data_text = '\n'.join(current_data_lines).strip()
        if data_text:
            blocks.append({
                "id": f"p{page_num}_b{block_idx}",
                "type": "text",
                "bbox": [0, 0, 0, 0],
                "text": data_text,
                "page_num": page_num,
                "source": "doc_parser",
                "score": 1.0,
                "meta": {"block_subtype": "data_rows"}
            })
            block_idx += 1
        current_data_lines.clear()

    for line in lines:
        stripped = line.strip()
        is_header = _is_doc_section_header(line)
        is_data = _is_data_row(line)
        # 목록 기호 시작 판정 (□, ◦, -, 1),   ◦ 등)
        is_list_item = len(stripped) < 200 and re.match(r'^\s*([□◦\-\u25e6\u2022\u25a1\u25cb\u2023\u25a0]|[\d]{1,2}[\.\)])', line)

        if is_header:
            _flush_data()
            _flush_body()
            blocks.append({
                "id": f"p{page_num}_b{block_idx}",
                "type": "title",
                "bbox": [0, 0, 0, 0],
                "text": line.strip(),
                "page_num": page_num,
                "source": "doc_parser",
                "score": 1.0,
                "meta": {}
            })
            block_idx += 1
        elif not stripped:
            # 빈 줄: 블록 단절 생성
            _flush_data()
            _flush_body()
        elif is_list_item:
            # 목록 아이템 시작: 이전 버퍼 비우고 새 블록 유도
            _flush_data()
            _flush_body()
            if is_data:
                current_data_lines.append(line)
            else:
                current_body_lines.append(line)
        elif is_data:
            # 데이터 행: 바디 버퍼가 있으면 비우고 데이터 버퍼로
            if current_body_lines:
                _flush_body()
            current_data_lines.append(line)
        else:
            # 일반 텍스트: 데이터 버퍼가 있으면 비우고 바디 버퍼로
            if current_data_lines:
                _flush_data()
            current_body_lines.append(line)

    _flush_data()
    _flush_body()
    
    # 흡수 최적화 (파편 문장 합치기)
    blocks = _absorb_tiny_fragments_sequential(blocks)
    
    return blocks

def _absorb_tiny_fragments_sequential(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """DOC/HWP 등 bbox가 없는 순차적 블록의 조각을 인접 블록으로 흡수"""
    if not blocks: return []
    out = []
    
    for i, b in enumerate(blocks):
        text = b.get("text", "")
        # Tiny chunk (under 30 chars, single line, no structural header numbers)
        if b["type"] == "text" and len(text) < 30 and "\n" not in text and not (len(text) < 5 and text.strip().isdigit()):
            # Absorb to the previous block if text type, otherwise next block
            if out and out[-1]["type"] == "text":
                out[-1]["text"] += "\n" + text.strip()
            elif i + 1 < len(blocks) and blocks[i+1]["type"] == "text":
                blocks[i+1]["text"] = text.strip() + "\n" + blocks[i+1].get("text", "")
            else:
                out.append(b)
        else:
            out.append(b)
            
    return out


def _generate_doc_rag_text(blocks: list[dict[str, Any]]) -> str:
    """DOC 전용 RAG 텍스트 생성 (PDF와 유사한 스키마 사용)"""
    parts = []
    for b in blocks:
        btype = b["type"]
        text = b.get("text", "").strip()
        if not text: continue
        
        if btype == "title":
            parts.append(f"\n[SECTION: {text}]\n")
        elif btype == "text":
            parts.append(text)
        elif btype == "table":
            parts.append(f"\n[TABLE data in text form: {text[:200]}...]\n")
    return "\n\n".join(parts).strip()


# ── public entry point ───────────────────────────────────────────────

def parse_doc(filepath: str) -> dict[str, Any]:
    ext = os.path.splitext(filepath)[1].lower()
    errors: list[str] = []

    strategies = []
    if ext == ".docx":
        strategies.append(("python-docx", _parse_docx))
    else:
        strategies.extend([
            ("tika", _parse_doc_tika),
            ("win32com", _parse_doc_win32),
            ("OLE binary", _parse_doc_binary),
        ])

    for name, fn in strategies:
        try:
            result = fn(filepath)
            result.setdefault("metadata", {})["parse_strategies_tried"] = (
                errors if errors else [f"{name} (first attempt)"]
            )
            
            # ── Document Routing ──────────────────────────────
            from .document_router import route_document
            pages = result.get("pages", [])
            meta = result.get("metadata", {})
            routing_result = route_document(None, pages, metadata=meta)
            
            doc_type = routing_result.get("document_type", "text_report")
            meta.update({
                "document_type": doc_type,
                "routing_signals": routing_result.get("routing_signals"),
                "routing_reason": routing_result.get("routing_reason")
            })
            
            for p in pages:
                p.setdefault("parser_debug", {})["document_type"] = doc_type
                # RAG Text Generation
                p["rag_text"] = _generate_doc_rag_text(p.get("blocks", []))
                
            return result
        except Exception as exc:
            msg = f"{name}: {exc}"
            errors.append(msg)
            logger.info("doc strategy '%s' failed: %s", name, exc)

    return {
        "pages": [],
        "metadata": {"parse_strategies_tried": errors},
        "status": "error",
        "error": "All parsing strategies failed: " + " | ".join(errors),
    }


# ── Strategy: python-docx (.docx) ───────────────────────────────────

def _parse_docx(filepath: str) -> dict[str, Any]:
    from docx import Document

    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = _normalise_text("\n".join(paragraphs))

    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    metadata: dict[str, Any] = {
        "page_count": 1,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "char_count": len(full_text),
        "parser_used": "python-docx",
    }

    sec_blocks = _split_into_section_blocks(full_text)
    pages = [{
        "page_num": 1,
        "page_width": 0,
        "page_height": 0,
        "preview_width": 0,
        "preview_height": 0,
        "preview_scale_x": 1.0,
        "preview_scale_y": 1.0,
        "coord_space": "page_points",
        "preview_image": None,
        "text": full_text,
        "tables": tables,
        "blocks": sec_blocks,
        "image_count": 0,
        "text_source": "native",
        "ocr_applied": False,
        "ocr_confidence": 0.0,
        "parser_debug": {
            "preview_generated": False,
            "preview_error": None,
            "native_text_chars": len(full_text),
            "ocr_used": False,
            "ocr_trigger_reason": "ocr_not_needed",
            "candidate_counts": {
                "raw_text_blocks": 1 if full_text else 0,
                "final_blocks": len(sec_blocks)
            },
            "block_type_counts": {
                "title": sum(1 for b in sec_blocks if b["type"] == "title"),
                "text": sum(1 for b in sec_blocks if b["type"] == "text")
            },
            "dropped_blocks": [],
            "bbox_warnings": [],
            "narrative_analysis": _analyze_narrative_characteristics(full_text)
        }
    }]
    
    metadata["narrative_analysis"] = pages[0]["parser_debug"]["narrative_analysis"]
    return {"pages": pages, "metadata": metadata, "status": "success"}


# ── Strategy 1: Tika ────────────────────────────────────────────────

def _parse_doc_tika(filepath: str) -> dict[str, Any]:
    try:
        from tika import parser as tika_parser
    except ImportError:
        raise RuntimeError("tika package not installed")

    parsed = tika_parser.from_file(filepath)
    content = parsed.get("content")
    if not content or not content.strip():
        raise ValueError("Tika returned empty content")

    full_text = _normalise_text(content)
    tika_meta = parsed.get("metadata") or {}

    metadata: dict[str, Any] = {
        "page_count": int(tika_meta.get("xmpTPg:NPages", 1) or 1),
        "char_count": len(full_text),
        "parser_used": "Apache Tika",
        "tika_content_type": tika_meta.get("Content-Type", ""),
    }

    sec_blocks = _split_into_section_blocks(full_text)
    pages = [{
        "page_num": 1,
        "page_width": 0,
        "page_height": 0,
        "preview_width": 0,
        "preview_height": 0,
        "preview_scale_x": 1.0,
        "preview_scale_y": 1.0,
        "coord_space": "page_points",
        "preview_image": None,
        "text": full_text,
        "tables": [],
        "blocks": sec_blocks,
        "image_count": 0,
        "text_source": "native",
        "ocr_applied": False,
        "ocr_confidence": 0.0,
        "parser_debug": {
            "preview_generated": False,
            "preview_error": None,
            "native_text_chars": len(full_text),
            "ocr_used": False,
            "ocr_trigger_reason": "ocr_not_needed",
            "candidate_counts": {
                "raw_text_blocks": 1 if full_text else 0,
                "final_blocks": len(sec_blocks)
            },
            "block_type_counts": {
                "title": sum(1 for b in sec_blocks if b["type"] == "title"),
                "text": sum(1 for b in sec_blocks if b["type"] == "text")
            },
            "dropped_blocks": [],
            "bbox_warnings": [],
            "narrative_analysis": _analyze_narrative_characteristics(full_text)
        }
    }]
    
    metadata["narrative_analysis"] = pages[0]["parser_debug"]["narrative_analysis"]
    return {"pages": pages, "metadata": metadata, "status": "success"}


# ── Strategy 2: pywin32 COM ─────────────────────────────────────────

def _parse_doc_win32(filepath: str) -> dict[str, Any]:
    try:
        import win32com.client  # type: ignore
    except ImportError:
        raise RuntimeError(
            "pywin32 is not installed. Install with: pip install pywin32  "
            "(requires Microsoft Word on this machine)"
        )

    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(filepath), ReadOnly=True)
        full_text = _normalise_text(doc.Content.Text)
        doc.Close(False)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass

    metadata: dict[str, Any] = {
        "page_count": 1,
        "char_count": len(full_text),
        "parser_used": "win32com (Microsoft Word)",
    }
    sec_blocks = _split_into_section_blocks(full_text)
    pages = [{
        "page_num": 1,
        "page_width": 0,
        "page_height": 0,
        "preview_width": 0,
        "preview_height": 0,
        "preview_scale_x": 1.0,
        "preview_scale_y": 1.0,
        "coord_space": "page_points",
        "preview_image": None,
        "text": full_text,
        "tables": [],
        "blocks": sec_blocks,
        "image_count": 0,
        "text_source": "native",
        "ocr_applied": False,
        "ocr_confidence": 0.0,
        "parser_debug": {
            "preview_generated": False,
            "preview_error": None,
            "native_text_chars": len(full_text),
            "ocr_used": False,
            "ocr_trigger_reason": "ocr_not_needed",
            "candidate_counts": {
                "raw_text_blocks": 1 if full_text else 0,
                "final_blocks": len(sec_blocks)
            },
            "block_type_counts": {
                "title": sum(1 for b in sec_blocks if b["type"] == "title"),
                "text": sum(1 for b in sec_blocks if b["type"] == "text")
            },
            "dropped_blocks": [],
            "bbox_warnings": [],
            "narrative_analysis": _analyze_narrative_characteristics(full_text)
        }
    }]
    
    metadata["narrative_analysis"] = pages[0]["parser_debug"]["narrative_analysis"]
    return {"pages": pages, "metadata": metadata, "status": "success"}


# ── Strategy 3: OLE binary piece-table ──────────────────────────────

def _parse_doc_binary(filepath: str) -> dict[str, Any]:
    import olefile

    if not olefile.isOleFile(filepath):
        raise ValueError("Not an OLE2 compound file")

    ole = olefile.OleFileIO(filepath)
    try:
        full_text = _extract_ole_text(ole)
    finally:
        ole.close()

    full_text = _normalise_text(full_text)

    metadata: dict[str, Any] = {
        "page_count": 1,
        "char_count": len(full_text),
        "parser_used": "OLE binary (piece table)",
    }
    sec_blocks = _split_into_section_blocks(full_text)
    pages = [{
        "page_num": 1,
        "page_width": 0,
        "page_height": 0,
        "preview_width": 0,
        "preview_height": 0,
        "preview_scale_x": 1.0,
        "preview_scale_y": 1.0,
        "coord_space": "page_points",
        "preview_image": None,
        "text": full_text,
        "tables": [],
        "blocks": sec_blocks,
        "image_count": 0,
        "text_source": "native",
        "ocr_applied": False,
        "ocr_confidence": 0.0,
        "parser_debug": {
            "preview_generated": False,
            "preview_error": None,
            "native_text_chars": len(full_text),
            "ocr_used": False,
            "ocr_trigger_reason": "ocr_not_needed",
            "candidate_counts": {
                "raw_text_blocks": 1 if full_text else 0,
                "final_blocks": len(sec_blocks)
            },
            "block_type_counts": {
                "title": sum(1 for b in sec_blocks if b["type"] == "title"),
                "text": sum(1 for b in sec_blocks if b["type"] == "text")
            },
            "dropped_blocks": [],
            "bbox_warnings": [],
            "narrative_analysis": _analyze_narrative_characteristics(full_text)
        }
    }]
    
    metadata["narrative_analysis"] = pages[0]["parser_debug"]["narrative_analysis"]
    return {"pages": pages, "metadata": metadata, "status": "success"}


def _extract_ole_text(ole: Any) -> str:
    word_doc = ole.openstream("WordDocument").read()

    magic = struct.unpack_from("<H", word_doc, 0)[0]
    if magic not in (0xA5EC, 0xA5DC):
        raise ValueError(f"Invalid Word magic: {hex(magic)}")

    flags = struct.unpack_from("<H", word_doc, 0x000A)[0]
    table_name = "1Table" if (flags & 0x0200) else "0Table"

    if not ole.exists(table_name):
        alt = "0Table" if table_name == "1Table" else "1Table"
        if ole.exists(alt):
            table_name = alt
        else:
            raise ValueError("No Table stream found")

    table_stream = ole.openstream(table_name).read()

    csw = struct.unpack_from("<H", word_doc, 0x20)[0]
    fibrg_w_end = 0x22 + csw * 2

    cslw = struct.unpack_from("<H", word_doc, fibrg_w_end)[0]
    fibrg_lw_start = fibrg_w_end + 2
    fibrg_lw_end = fibrg_lw_start + cslw * 4

    ccpText = struct.unpack_from("<i", word_doc, fibrg_lw_start)[0]

    cb_rg = struct.unpack_from("<H", word_doc, fibrg_lw_end)[0]
    fclcb_start = fibrg_lw_end + 2

    if cb_rg < 34:
        raise ValueError(f"FIBRgFcLcb too small ({cb_rg} pairs)")

    fc_clx = struct.unpack_from("<I", word_doc, fclcb_start + 33 * 8)[0]
    lcb_clx = struct.unpack_from("<I", word_doc, fclcb_start + 33 * 8 + 4)[0]

    if lcb_clx == 0:
        raise ValueError("CLX has zero length")

    clx = table_stream[fc_clx: fc_clx + lcb_clx]

    off = 0
    while off < len(clx) and clx[off] == 0x01:
        cb = struct.unpack_from("<H", clx, off + 1)[0]
        off += 3 + cb

    if off >= len(clx) or clx[off] != 0x02:
        raise ValueError("Pcdt marker (0x02) not found")

    off += 1
    lcb_pcdt = struct.unpack_from("<I", clx, off)[0]
    off += 4
    plc = clx[off: off + lcb_pcdt]

    n_pieces = (lcb_pcdt - 4) // 12
    if n_pieces <= 0:
        raise ValueError("No text pieces found")

    cps = [struct.unpack_from("<I", plc, i * 4)[0] for i in range(n_pieces + 1)]
    pcd_base = (n_pieces + 1) * 4

    parts: list[str] = []
    for i in range(n_pieces):
        cp_start, cp_end = cps[i], cps[i + 1]
        if cp_start >= ccpText:
            break
        n_chars = min(cp_end, ccpText) - cp_start

        pcd = plc[pcd_base + i * 8: pcd_base + (i + 1) * 8]
        fc_raw = struct.unpack_from("<I", pcd, 2)[0]
        compressed = bool(fc_raw & (1 << 30))
        fc_val = fc_raw & 0x3FFFFFFF

        if compressed:
            boff = fc_val // 2
            parts.append(word_doc[boff: boff + n_chars].decode("cp1252", errors="replace"))
        else:
            boff = fc_val
            parts.append(word_doc[boff: boff + n_chars * 2].decode("utf-16-le", errors="replace"))

    return "".join(parts)
